import os
import re
import sys
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

if sys.stdout.encoding != 'utf-8':
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except Exception:
        pass

def get_safe_image_path(img_path):
    if not img_path or not os.path.exists(img_path):
        return None
    try:
        from PIL import Image
        img = Image.open(img_path)
        if img.format not in ['PNG', 'JPEG', 'JPG']:
            dir_name = os.path.dirname(img_path)
            base_name = os.path.basename(img_path)
            name_part, _ = os.path.splitext(base_name)
            safe_path = os.path.join(dir_name, f"fixed_{name_part}.png")
            img.convert('RGB').save(safe_path, 'PNG')
            print(f"Auto-converted {img_path} ({img.format}) -> {safe_path}")
            return safe_path
        return img_path
    except Exception as e:
        print(f"Error checking image {img_path}: {e}")
        return img_path

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def process_inline_markdown(p, raw_text):
    pattern = re.compile(r'(\*\*|__)(.*?)\1|(`)(.*?)\3|(\*)(.*?)\5')
    last_end = 0
    for match in pattern.finditer(raw_text):
        start, end = match.span()
        if start > last_end:
            normal_txt = raw_text[last_end:start]
            if normal_txt:
                r = p.add_run(normal_txt)
                r.font.name = 'Times New Roman'
                
        if match.group(1): # Bold
            styled_txt = match.group(2)
            r = p.add_run(styled_txt)
            r.bold = True
            r.font.name = 'Times New Roman'
        elif match.group(3): # Code
            styled_txt = match.group(4)
            r = p.add_run(styled_txt)
            r.font.name = 'Courier New'
            r.font.size = Pt(11)
        elif match.group(5): # Italic
            styled_txt = match.group(6)
            r = p.add_run(styled_txt)
            r.italic = True
            r.font.name = 'Times New Roman'
            
        last_end = end
        
    if last_end < len(raw_text):
        rem_txt = raw_text[last_end:]
        if rem_txt:
            r = p.add_run(rem_txt)
            r.font.name = 'Times New Roman'

def draw_word_table(doc, rows_data):
    if not rows_data:
        return
    rows_count = len(rows_data)
    cols_count = len(rows_data[0])
    
    table = doc.add_table(rows=rows_count, cols=cols_count)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for r_idx, row in enumerate(rows_data):
        for c_idx in range(min(cols_count, len(row))):
            cell = table.cell(r_idx, c_idx)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            
            val = row[c_idx]
            if r_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(val)
                run.bold = True
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(val)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            
    doc.add_paragraph()

def parse_markdown_to_docx(doc, filepath):
    if not os.path.exists(filepath):
        print(f"File not found: {filepath}")
        return
    print(f"Injecting file: {filepath}")
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    idx = 0
    in_table = False
    table_rows = []
    
    while idx < len(lines):
        line = lines[idx].strip()
        
        if not line:
            if in_table:
                in_table = False
                draw_word_table(doc, table_rows)
                table_rows = []
            idx += 1
            continue
            
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(line[2:])
            r.bold = True
            r.font.size = Pt(13)
            r.font.name = 'Times New Roman'
        elif line.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(line[3:])
            r.bold = True
            r.font.size = Pt(13)
            r.font.name = 'Times New Roman'
        elif line.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(line[4:])
            r.bold = True
            r.italic = True
            r.font.size = Pt(13)
            r.font.name = 'Times New Roman'
        elif line.startswith('#### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(line[5:])
            r.bold = True
            r.font.size = Pt(13)
            r.font.name = 'Times New Roman'
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_after = Pt(4)
            p.add_run('•  ').font.name = 'Times New Roman'
            process_inline_markdown(p, line[2:])
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_after = Pt(4)
            match = re.match(r'^(\d+)\.\s(.*)', line)
            p.add_run(f'{match.group(1)}.  ').font.name = 'Times New Roman'
            process_inline_markdown(p, match.group(2))
        elif line.startswith('|'):
            in_table = True
            parts = [p.strip() for p in line.split('|')[1:-1]]
            if all(re.match(r'^:?-+:?$', p) for p in parts):
                idx += 1
                continue
            table_rows.append(parts)
        else:
            if in_table:
                in_table = False
                draw_word_table(doc, table_rows)
                table_rows = []
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.first_line_indent = Cm(1.25)
            p.paragraph_format.space_after = Pt(6)
            process_inline_markdown(p, line)
            
        idx += 1
        
    if in_table and table_rows:
        draw_word_table(doc, table_rows)

def parse_txt_to_docx(doc, filepath):
    if not os.path.exists(filepath):
        print(f"File not found: {filepath}")
        return
    print(f"Injecting text file: {filepath}")
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
        
    paragraphs = content.split('\n\n')
    for p_text in paragraphs:
        p_text = p_text.strip()
        if not p_text:
            continue
            
        is_heading = False
        if len(p_text) < 80:
            if p_text.startswith('5.') or p_text.startswith('Chương') or any(p_text.startswith(x) for x in ['Giao diện', 'Hệ thống', 'Tính năng', 'Trợ lý', 'Quản trị', 'Bảo mật', 'Chưa', 'Thiếu', 'Hiệu suất', 'Yêu cầu', 'Hạn chế']):
                is_heading = True
                
        p = doc.add_paragraph()
        if is_heading:
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(p_text)
            r.bold = True
            r.font.size = Pt(13)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.first_line_indent = Cm(1.25)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(p_text)
            
        r.font.name = 'Times New Roman'

def create_report():
    doc = Document()
    
    # Configure styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    
    # Page setup (margins)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
    
    # Helper to add page borders (native XML PG borders)
    def add_page_borders(section):
        sectPr = section._sectPr
        # Remove existing if any
        existing_borders = sectPr.find(qn('w:pgBorders'))
        if existing_borders is not None:
            sectPr.remove(existing_borders)
            
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        
        # 4 borders: top, left, bottom, right
        # We put a nice frame with size 12 (1.5 pt) and 24pt spacing
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '12')
            border.set(qn('w:space'), '24')
            border.set(qn('w:color'), '000000')
            pgBorders.append(border)
        sectPr.append(pgBorders)

    # Helper to remove borders
    def remove_page_borders(section):
        sectPr = section._sectPr
        existing_borders = sectPr.find(qn('w:pgBorders'))
        if existing_borders is not None:
            sectPr.remove(existing_borders)

    # ====== PAGE 1: PHỤ LỤC HƯỚNG DẪN ======
    p_guidelines = doc.add_paragraph()
    p_guidelines.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # BM5 Box at the top right
    # Let's make a tiny table for BM5 box
    bm_table = doc.add_table(rows=1, cols=1)
    bm_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    cell = bm_table.cell(0, 0)
    cell.width = Cm(1.5)
    p_bm = cell.paragraphs[0]
    p_bm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_bm = p_bm.add_run('BM5')
    r_bm.font.name = 'Times New Roman'
    r_bm.font.size = Pt(11)
    
    # Spacing
    for _ in range(8):
        doc.add_paragraph()
        
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run('CÁC PHỤ LỤC HƯỚNG DẪN LÀM ĐỒ ÁN,\nKHÓA LUẬN TỐT NGHIỆP')
    r_title.bold = True
    r_title.font.size = Pt(16)
    r_title.font.name = 'Times New Roman'
    
    doc.add_page_break()

    # ====== PAGE 2: BÌA CHÍNH (BÌA CỨNG) ======
    section_cover = doc.sections[-1]
    add_page_borders(section_cover)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('(Mẫu trang bìa – in trên bìa cứng)')
    r.italic = True
    r.font.size = Pt(11)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('ĐẠI HỌC TRÀ VINH')
    r.bold = True
    r.font.size = Pt(16)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run('TRƯỜNG KỸ THUẬT VÀ CÔNG NGHỆ')
    r.bold = True
    r.font.size = Pt(16)
    
    # Logo
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_path = get_safe_image_path(os.path.join('images', 'logo_tvu.png'))
    if logo_path:
        p_logo.add_run().add_picture(logo_path, width=Cm(3.0))
    else:
        # Fallback to general logo path
        p_logo.add_run('[LOGO ĐẠI HỌC TRÀ VINH]').bold = True
        
    p_logo.paragraph_format.space_after = Pt(36)
    p_logo.paragraph_format.space_before = Pt(24)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('ĐỒ ÁN TỐT NGHIỆP')
    r.bold = True
    r.font.size = Pt(16)
    p.paragraph_format.space_after = Pt(18)

    p_topic = doc.add_paragraph()
    p_topic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_topic.paragraph_format.space_after = Pt(60)
    p_topic.paragraph_format.line_spacing = 1.3
    r = p_topic.add_run('Xây dựng Website kết nối sinh viên ngành y với người có nhu cầu chăm sóc sức khỏe tại nhà')
    r.bold = True
    r.font.size = Pt(18) # standard 18-30 for title

    # Info table
    info_table = doc.add_table(rows=5, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Configure column widths
    for row in info_table.rows:
        row.cells[0].width = Cm(6.0)
        row.cells[1].width = Cm(8.0)
        
    labels = [
        'Giảng viên hướng dẫn:',
        'Sinh viên thực hiện:',
        'Mã số sinh viên:',
        'Lớp:',
        'Khóa:'
    ]
    values = [
        'ThS NGUYỄN HOÀNG DUY THIÊN',
        'TRẦM TẤN KHA',
        '110122087',
        'DA22TTD',
        '2022'
    ]
    
    for i in range(5):
        cell_lbl = info_table.cell(i, 0)
        cell_val = info_table.cell(i, 1)
        
        p_lbl = cell_lbl.paragraphs[0]
        r_lbl = p_lbl.add_run(labels[i])
        r_lbl.font.size = Pt(12)
        r_lbl.font.name = 'Times New Roman'
        
        p_val = cell_val.paragraphs[0]
        r_val = p_val.add_run(values[i])
        r_val.bold = True
        r_val.font.size = Pt(12)
        r_val.font.name = 'Times New Roman'
        
    # Clear borders of info table
    for row in info_table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for b in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{b}')
                border.set(qn('w:val'), 'none')
                tcBorders.append(border)
            tcPr.append(tcBorders)

    # Date at the bottom
    for _ in range(5):
        doc.add_paragraph()
        
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_date = p_date.add_run('Vĩnh Long, tháng ... năm ...')
    r_date.bold = True
    r_date.font.size = Pt(13)

    doc.add_page_break()

    # ====== PAGE 3: BÌA PHỤ (TRANG LÓT) ======
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('(Mẫu trang lót – in giấy thường)')
    r.italic = True
    r.font.size = Pt(11)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('ĐẠI HỌC TRÀ VINH')
    r.bold = True
    r.font.size = Pt(16)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run('TRƯỜNG KỸ THUẬT VÀ CÔNG NGHỆ')
    r.bold = True
    r.font.size = Pt(16)
    
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if logo_path:
        p_logo.add_run().add_picture(logo_path, width=Cm(3.0))
    else:
        p_logo.add_run('[LOGO ĐẠI HỌC TRÀ VINH]').bold = True
    p_logo.paragraph_format.space_after = Pt(36)
    p_logo.paragraph_format.space_before = Pt(24)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('ĐỒ ÁN TỐT NGHIỆP')
    r.bold = True
    r.font.size = Pt(16)
    p.paragraph_format.space_after = Pt(18)

    # In the template page 3 has:
    # ................................................................
    # ................................................................
    # (Bold, size 18-30, tùy theo số chữ... của tên đề tài)
    # However, to be a complete, beautiful report, we write the actual topic name inside a dotted structure or directly. Let's write the actual topic.
    p_topic2 = doc.add_paragraph()
    p_topic2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_topic2.paragraph_format.space_after = Pt(60)
    p_topic2.paragraph_format.line_spacing = 1.3
    r = p_topic2.add_run('Xây dựng Website kết nối sinh viên ngành y với người có nhu cầu chăm sóc sức khỏe tại nhà')
    r.bold = True
    r.font.size = Pt(18)

    # Info table lót
    info_table2 = doc.add_table(rows=5, cols=2)
    info_table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row in info_table2.rows:
        row.cells[0].width = Cm(6.0)
        row.cells[1].width = Cm(8.0)
        
    for i in range(5):
        cell_lbl = info_table2.cell(i, 0)
        cell_val = info_table2.cell(i, 1)
        
        p_lbl = cell_lbl.paragraphs[0]
        r_lbl = p_lbl.add_run(labels[i])
        r_lbl.font.size = Pt(12)
        r_lbl.font.name = 'Times New Roman'
        
        p_val = cell_val.paragraphs[0]
        r_val = p_val.add_run(values[i])
        r_val.bold = True
        r_val.font.size = Pt(12)
        r_val.font.name = 'Times New Roman'
        
    for row in info_table2.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for b in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{b}')
                border.set(qn('w:val'), 'none')
                tcBorders.append(border)
            tcPr.append(tcBorders)

    # Date at the bottom lót
    for _ in range(5):
        doc.add_paragraph()
        
    p_date2 = doc.add_paragraph()
    p_date2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_date2 = p_date2.add_run('Vĩnh Long, tháng ... năm ...')
    r_date2.bold = True
    r_date2.font.size = Pt(13)

    doc.add_page_break()
    
    # We create a new section for the body, so we can remove page borders
    new_section = doc.add_section()
    remove_page_borders(new_section)

    # Configure page numbers in footer
    footer = new_section.footer
    footer.is_linked_to_previous = False
    p_footer = footer.paragraphs[0]
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_footer.text = ""
    run_footer = p_footer.add_run()
    run_footer.font.name = 'Times New Roman'
    run_footer.font.size = Pt(11)
    add_page_number(run_footer)

    # ====== PAGE 4: MỤC LỤC ======
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run('MỤC LỤC')
    r.bold = True
    r.font.size = Pt(14)
    
    # Dotted table of contents
    def add_toc_line(title_text, page_str, bold=False, level=0):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(4)
        
        # Left indent based on hierarchy level
        if level == 1:
            p.paragraph_format.left_indent = Cm(0.6)
        elif level == 2:
            p.paragraph_format.left_indent = Cm(1.2)
        elif level == 3:
            p.paragraph_format.left_indent = Cm(1.8)
            
        r_title = p.add_run(title_text)
        r_title.bold = bold
        r_title.font.name = 'Times New Roman'
        r_title.font.size = Pt(13 if level == 0 else 12)
        
        # Dot leaders
        p_width_cm = 16.0 - (level * 0.6)
        # Use tab stop for alignment
        tab_stops = p.paragraph_format.tab_stops
        tab_stop = tab_stops.add_tab_stop(Cm(16.0 - (level * 0.6)), alignment=1) # 1 = right align
        
        # Add tab and page number
        p.add_run('\t')
        r_page = p.add_run(page_str)
        r_page.bold = bold
        r_page.font.name = 'Times New Roman'
        r_page.font.size = Pt(13 if level == 0 else 12)
        
    add_toc_line('Lời mở đầu', '5', bold=True)
    add_toc_line('Lời cảm ơn', '6', bold=True)
    add_toc_line('Nhận xét của Cơ quan thực tập', '7', bold=True)
    add_toc_line('Nhận xét của Giảng viên hướng dẫn', '8', bold=True)
    add_toc_line('Bản nhận xét Đồ án tốt nghiệp của Giảng viên hướng dẫn', '9', bold=True)
    add_toc_line('Nhận xét của Giảng viên chấm', '11', bold=True)
    add_toc_line('Bản nhận xét Đồ án tốt nghiệp của Cán bộ chấm', '12', bold=True)
    add_toc_line('Danh mục các bảng, sơ đồ, hình', '14', bold=True)
    add_toc_line('Kí hiệu các cụm từ viết tắt', '15', bold=True)
    
    add_toc_line('TÓM TẮT ĐỒ ÁN', '16', bold=True)
    add_toc_line('MỞ ĐẦU', '17', bold=True)
    
    add_toc_line('CHƯƠNG 1: TỔNG QUAN', '18', bold=True)
    add_toc_line('1.1. Giới thiệu đề tài', '18', level=1)
    add_toc_line('1.2. Tình hình nghiên cứu hiện tại', '18', level=1)
    add_toc_line('1.3. Mục tiêu của đề tài', '19', level=1)
    add_toc_line('1.4. Ý nghĩa khoa học và thực tiễn', '19', level=1)
    
    add_toc_line('CHƯƠNG 2: NGHIÊN CỨU LÝ THUYẾT', '21', bold=True)
    add_toc_line('2.1. Tổng quan về phát triển Web', '21', level=1)
    add_toc_line('2.2. Ngôn ngữ lập trình PHP', '21', level=1)
    add_toc_line('2.3. Hệ quản trị cơ sở dữ liệu MySQL', '22', level=1)
    add_toc_line('2.4. Công nghệ Front-end', '22', level=1)
    add_toc_line('2.5. Môi trường phát triển XAMPP', '23', level=1)
    add_toc_line('2.6. Phương pháp phân tích thiết kế hệ thống', '23', level=1)
    
    add_toc_line('CHƯƠNG 3: HIỆN THỰC HÓA NGHIÊN CỨU', '25', bold=True)
    add_toc_line('3.1. Phân tích yêu cầu hệ thống', '25', level=1)
    add_toc_line('3.2. Thiết kế hệ thống', '26', level=1)
    add_toc_line('3.3. Cài đặt hệ thống', '27', level=1)
    
    add_toc_line('CHƯƠNG 4: KẾT QUẢ NGHIÊN CỨU', '30', bold=True)
    add_toc_line('4.1. Giao diện người dùng', '30', level=1)
    add_toc_line('4.2. Đánh giá kết quả', '36', level=1)
    
    add_toc_line('CHƯƠNG 5: KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN', '38', bold=True)
    add_toc_line('5.1. Kết luận', '38', level=1)
    add_toc_line('5.2. Hướng phát triển', '39', level=1)
    
    add_toc_line('DANH MỤC TÀI LIỆU THAM KHẢO', '41', bold=True)
    add_toc_line('PHỤ LỤC', '42', bold=True)

    # Footnote/Note matching template
    for _ in range(3):
        doc.add_paragraph()
    p_note = doc.add_paragraph()
    p_note.paragraph_format.space_before = Pt(12)
    p_note.paragraph_format.line_spacing = 1.15
    r = p_note.add_run('Ghi chú:\n')
    r.bold = True
    r.italic = True
    r.font.size = Pt(10)
    r = p_note.add_run('- In đậm và in hoa tiêu đề của các chương, mục lớn\n'
                       '- Chữ số thứ nhất chỉ thứ tự chương\n'
                       '- Chữ số thứ 2 chỉ thứ tự mục trong chương\n'
                       '- Chữ số thứ 3,... chỉ thứ tự các tiểu mục\n'
                       '- Các tiểu mục trình bày trong mục lục không quá 04 cấp')
    r.italic = True
    r.font.size = Pt(10)

    doc.add_page_break()

    # ====== PAGE 5: LỜI MỞ ĐẦU (MẪU) ======
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('(Mẫu)\n')
    r.italic = True
    r.font.size = Pt(10)
    r = p.add_run('LỜI MỞ ĐẦU\n')
    r.bold = True
    r.font.size = Pt(14)
    r = p.add_run('(Bold, size 14, xếp sau trang lót)')
    r.italic = True
    r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(24)
    
    # Border box for the outline of text (as seen in image page 5)
    # We can create a 1x1 table to represent the bordered frame in the template
    border_table = doc.add_table(rows=1, cols=1)
    border_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = border_table.cell(0, 0)
    p_cell = cell.paragraphs[0]
    p_cell.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    r_body = p_cell.add_run('size 13, ........................................................................................\n'
                            '..........................................................................................................\n'
                            '..........................................................................................................\n'
                            '..........................................................................................................\n'
                            '..........................................................................................................\n'
                            '..........................................................................................................\n'
                            '..........................................................................................................\n'
                            '..........................................................................................................')
    r_body.font.size = Pt(13)
    r_body.font.name = 'Times New Roman'
    p_cell.paragraph_format.line_spacing = 1.5
    
    doc.add_page_break()

    # ====== PAGE 6: LỜI CẢM ƠN (MẪU) ======
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('(Mẫu)\n')
    r.italic = True
    r.font.size = Pt(10)
    r = p.add_run('LỜI CẢM ƠN\n')
    r.bold = True
    r.font.size = Pt(14)
    r = p.add_run('(bold, size 14, xếp sau trang lời mở đầu)')
    r.italic = True
    r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(24)

    border_table = doc.add_table(rows=1, cols=1)
    border_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = border_table.cell(0, 0)
    p_cell = cell.paragraphs[0]
    p_cell.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_body = p_cell.add_run('size 13, ........................................................................................\n'
                            '..........................................................................................................\n'
                            '..........................................................................................................\n'
                            '..........................................................................................................\n'
                            '..........................................................................................................\n'
                            '..........................................................................................................\n'
                            '..........................................................................................................\n'
                            '..........................................................................................................')
    r_body.font.size = Pt(13)
    r_body.font.name = 'Times New Roman'
    p_cell.paragraph_format.line_spacing = 1.5

    doc.add_page_break()

    # ====== PAGE 7: NHẬN XÉT CƠ QUAN THỰC TẬP ======
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('(Mẫu)\n')
    r.italic = True
    r.font.size = Pt(10)
    r = p.add_run('NHẬN XÉT\n')
    r.bold = True
    r.font.size = Pt(14)
    r = p.add_run('(Của cơ quan thực tập, nếu có)\nBold, size 14, xếp sau trang Lời cảm ơn')
    r.italic = True
    r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(24)
    
    p_rx = doc.add_paragraph()
    p_rx.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p_rx.add_run('Size 13, ..................................................................................................................\n'
                     '....................................................................................................................................\n'
                     '....................................................................................................................................\n'
                     '....................................................................................................................................\n'
                     '....................................................................................................................................\n'
                     '....................................................................................................................................\n'
                     '....................................................................................................................................\n'
                     '....................................................................................................................................\n'
                     '....................................................................................................................................\n'
                     '....................................................................................................................................\n'
                     '....................................................................................................................................\n'
                     '....................................................................................................................................')
    r.font.size = Pt(13)
    r.font.name = 'Times New Roman'
    p_rx.paragraph_format.line_spacing = 1.5
    
    doc.add_page_break()

    # ====== PAGE 8: NHẬN XÉT CỦA GIẢNG VIÊN HƯỚNG DẪN ======
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('(Mẫu)\n')
    r.italic = True
    r.font.size = Pt(10)
    r = p.add_run('NHẬN XÉT\n')
    r.bold = True
    r.font.size = Pt(14)
    r = p.add_run('(Của giảng viên hướng dẫn trong đồ án, khoá luận của sinh viên)\n(Bold, size 14, xếp sau trang nhận xét của cơ quan thực tập)')
    r.italic = True
    r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(24)

    p_rx = doc.add_paragraph()
    p_rx.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p_rx.add_run('Size 13, ..................................................................................................................\n'
                     '....................................................................................................................................\n'
                     '....................................................................................................................................\n'
                     '....................................................................................................................................\n'
                     '....................................................................................................................................\n'
                     '....................................................................................................................................\n'
                     '....................................................................................................................................\n'
                     '....................................................................................................................................\n'
                     '....................................................................................................................................\n'
                     '....................................................................................................................................\n'
                     '....................................................................................................................................\n'
                     '....................................................................................................................................')
    r.font.size = Pt(13)
    r.font.name = 'Times New Roman'
    p_rx.paragraph_format.line_spacing = 1.5

    for _ in range(2):
        doc.add_paragraph()
    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p_sig.add_run('Giảng viên hướng dẫn\n(ký và ghi rõ họ tên)')
    r.bold = True
    r.font.size = Pt(13)
    
    doc.add_page_break()

    # ====== PAGE 9 & 10: BẢN NHẬN XÉT CỦA GIẢNG VIÊN HƯỚNG DẪN ======
    p_hdr = doc.add_table(rows=1, cols=2)
    p_hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
    p_hdr.rows[0].cells[0].width = Cm(7.0)
    p_hdr.rows[0].cells[1].width = Cm(8.0)
    
    cell_l = p_hdr.cell(0, 0)
    p_l = cell_l.paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_l.add_run('UBND TỈNH VĨNH LONG\n')
    r.font.size = Pt(11)
    r = p_l.add_run('ĐẠI HỌC TRÀ VINH')
    r.bold = True
    r.font.size = Pt(11)
    
    cell_r = p_hdr.cell(0, 1)
    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_r.add_run('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n')
    r.bold = True
    r.font.size = Pt(11)
    r = p_r.add_run('Độc lập - Tự do - Hạnh phúc')
    r.italic = True
    r.font.size = Pt(11)
    
    for row in p_hdr.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for b in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{b}')
                border.set(qn('w:val'), 'none')
                tcBorders.append(border)
            tcPr.append(tcBorders)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run('BẢN NHẬN XÉT ĐỒ ÁN, KHÓA LUẬN TỐT NGHIỆP\n')
    r.bold = True
    r.font.size = Pt(14)
    r = p.add_run('(Của giảng viên hướng dẫn)')
    r.italic = True
    r.font.size = Pt(12)

    # Info fields
    def add_dotted_field(label, value_len=80):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(4)
        r = p.add_run(label)
        r.font.size = Pt(12)
        r_dots = p.add_run('.' * value_len)
        r_dots.font.size = Pt(12)
        
    add_dotted_field('Họ và tên sinh viên: ', 65)
    add_dotted_field('Ngành: ', 85)
    add_dotted_field('MSSV: ', 85)
    add_dotted_field('Khóa: ', 85)
    add_dotted_field('Tên đề tài: ', 80)
    add_dotted_field('Họ và tên Giáo viên hướng dẫn: ', 50)
    add_dotted_field('Chức danh: ', 50)
    add_dotted_field('Học vị: ', 80)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run('NHẬN XÉT')
    r.bold = True
    r.font.size = Pt(13)

    def add_numbered_section(number_str, title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f'{number_str}. {title}:')
        r.bold = True
        r.font.size = Pt(12)
        
        p_dots = doc.add_paragraph()
        p_dots.paragraph_format.space_after = Pt(6)
        r_dots = p_dots.add_run('....................................................................................................................................\n'
                               '....................................................................................................................................\n'
                               '....................................................................................................................................')
        r_dots.font.size = Pt(12)

    add_numbered_section('1', 'Nội dung đề tài')
    add_numbered_section('2', 'Ưu điểm')
    add_numbered_section('3', 'Khuyết điểm')

    doc.add_page_break() # Goes to Page 10

    add_numbered_section('4', 'Điểm mới của đề tài')
    add_numbered_section('5', 'Giá trị thực tiễn của đề tài')
    add_numbered_section('7', 'Đề nghị sửa chữa bổ sung')
    add_numbered_section('8', 'Đánh giá')

    for _ in range(2):
        doc.add_paragraph()
    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p_sig.add_run('Vĩnh Long, ngày ... tháng ... năm ...\n')
    r.italic = True
    r.font.size = Pt(12)
    r = p_sig.add_run('Giảng viên hướng dẫn\n')
    r.bold = True
    r.font.size = Pt(13)
    r = p_sig.add_run('(Ký & ghi rõ họ tên)')
    r.italic = True
    r.font.size = Pt(11)

    doc.add_page_break()

    # ====== PAGE 11: NHẬN XÉT CỦA GIẢNG VIÊN CHẤM ======
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('(Mẫu)\n')
    r.italic = True
    r.font.size = Pt(10)
    r = p.add_run('NHẬN XÉT\n')
    r.bold = True
    r.font.size = Pt(14)
    r = p.add_run('(Của giảng viên chấm trong đồ án, khoá luận của sinh viên)\n(Bold, size 14, xếp sau trang Nhận xét của giảng viên hướng dẫn)')
    r.italic = True
    r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(24)

    p_rx = doc.add_paragraph()
    p_rx.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p_rx.add_run('size 13, ..................................................................................................................\n'
                     '....................................................................................................................................\n'
                     '....................................................................................................................................\n'
                     '....................................................................................................................................\n'
                     '....................................................................................................................................\n'
                     '....................................................................................................................................\n'
                     '....................................................................................................................................\n'
                     '....................................................................................................................................\n'
                     '....................................................................................................................................\n'
                     '....................................................................................................................................\n'
                     '....................................................................................................................................\n'
                     '....................................................................................................................................')
    r.font.size = Pt(13)
    r.font.name = 'Times New Roman'
    p_rx.paragraph_format.line_spacing = 1.5

    for _ in range(2):
        doc.add_paragraph()
    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p_sig.add_run('Giảng viên chấm\n(ký và ghi rõ họ tên)')
    r.bold = True
    r.font.size = Pt(13)

    doc.add_page_break()

    # ====== PAGE 12 & 13: BẢN NHẬN XÉT CỦA CÁN BỘ CHẤM ======
    p_hdr = doc.add_table(rows=1, cols=2)
    p_hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
    p_hdr.rows[0].cells[0].width = Cm(7.0)
    p_hdr.rows[0].cells[1].width = Cm(8.0)
    
    cell_l = p_hdr.cell(0, 0)
    p_l = cell_l.paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_l.add_run('UBND TỈNH VĨNH LONG\n')
    r.font.size = Pt(11)
    r = p_l.add_run('ĐẠI HỌC TRÀ VINH')
    r.bold = True
    r.font.size = Pt(11)
    
    cell_r = p_hdr.cell(0, 1)
    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_r.add_run('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n')
    r.bold = True
    r.font.size = Pt(11)
    r = p_r.add_run('Độc lập - Tự do - Hạnh phúc')
    r.italic = True
    r.font.size = Pt(11)
    
    for row in p_hdr.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for b in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{b}')
                border.set(qn('w:val'), 'none')
                tcBorders.append(border)
            tcPr.append(tcBorders)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run('BẢN NHẬN XÉT ĐỒ ÁN, KHÓA LUẬN TỐT NGHIỆP\n')
    r.bold = True
    r.font.size = Pt(14)
    r = p.add_run('(Của cán bộ chấm đồ án, khóa luận)')
    r.italic = True
    r.font.size = Pt(12)

    add_dotted_field('Họ và tên người nhận xét: ', 55)
    add_dotted_field('Chức danh: ', 80)
    add_dotted_field('Học vị: ', 80)
    add_dotted_field('Chuyên ngành: ', 80)
    add_dotted_field('Cơ quan công tác: ', 65)
    add_dotted_field('Tên sinh viên: ', 65)
    add_dotted_field('Tên đề tài đồ án, khóa luận tốt nghiệp: ', 45)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run('I. Ý KIẾN NHẬN XÉT')
    r.bold = True
    r.font.size = Pt(13)

    add_numbered_section('1', 'Nội dung')
    add_numbered_section('2', 'Điểm mới của các kết quả của đồ án, khóa luận')
    add_numbered_section('3', 'Ứng dụng thực tế')

    doc.add_page_break() # Page 13

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run('II. CÁC VẤN ĐỀ CẦN LÀM RÕ\n')
    r.bold = True
    r.font.size = Pt(13)
    r = p.add_run('(Các câu hỏi của giáo viên phản biện)')
    r.italic = True
    r.font.size = Pt(11)
    
    p_dots = doc.add_paragraph()
    p_dots.paragraph_format.space_after = Pt(6)
    r_dots = p_dots.add_run('....................................................................................................................................\n'
                           '....................................................................................................................................\n'
                           '....................................................................................................................................\n'
                           '....................................................................................................................................')
    r_dots.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run('III. KẾT LUẬN\n')
    r.bold = True
    r.font.size = Pt(13)
    r = p.add_run('(Ghi rõ đồng ý hay không đồng ý cho bảo vệ đồ án khóa luận tốt nghiệp)')
    r.italic = True
    r.font.size = Pt(11)

    p_dots = doc.add_paragraph()
    p_dots.paragraph_format.space_after = Pt(6)
    r_dots = p_dots.add_run('....................................................................................................................................\n'
                           '....................................................................................................................................\n'
                           '....................................................................................................................................')
    r_dots.font.size = Pt(12)

    for _ in range(2):
        doc.add_paragraph()
    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p_sig.add_run('....................., ngày ... tháng ... năm ...\n')
    r.italic = True
    r.font.size = Pt(12)
    r = p_sig.add_run('Người nhận xét\n')
    r.bold = True
    r.font.size = Pt(13)
    r = p_sig.add_run('(Ký & ghi rõ họ tên)')
    r.italic = True
    r.font.size = Pt(11)

    doc.add_page_break()

    # ====== PAGE 14: DANH MỤC CÁC BẢNG, SƠ ĐỒ, HÌNH ======
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('(Mẫu)\n')
    r.italic = True
    r.font.size = Pt(10)
    r = p.add_run('DANH MỤC CÁC BẢNG, SƠ ĐỒ, HÌNH\n')
    r.bold = True
    r.font.size = Pt(14)
    r = p.add_run('(In đậm, in hoa, size 14)')
    r.italic = True
    r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(24)

    # Dotted table
    def add_list_item(label, page):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run(label)
        r.font.size = Pt(13)
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(15.5), alignment=1)
        p.add_run('\t')
        r_page = p.add_run(page)
        r_page.font.size = Pt(13)

    add_list_item('BANG 1.1: Bảng đánh giá mức độ hài lòng về dịch vụ y tế tại nhà', '20')
    add_list_item('BANG 3.1: Bảng Users - Thông tin người dùng', '26')
    add_list_item('BANG 3.2: Bảng Posts - Bài đăng tuyển dụng & ứng tuyển', '26')
    add_list_item('BANG 3.3: Bảng Messages - Tin nhắn trao đổi', '27')
    add_list_item('BANG 4.1: Bảng đánh giá các chức năng đã hoàn thành', '36')
    
    add_list_item('SƠ ĐỒ 1.1: Biểu đồ luồng hoạt động kết nối y tế', '19')
    add_list_item('SƠ ĐỒ 3.1: Biểu đồ Use Case tổng quan hệ thống', '25')
    
    add_list_item('HÌNH 3.1: Biểu đồ ERD - Quan hệ thực thể cơ sở dữ liệu', '26')
    add_list_item('HÌNH 4.1: Giao diện trang chủ tìm kiếm tin đăng', '30')
    add_list_item('HÌNH 4.2: Giao diện đăng ký tài khoản (Phân quyền email)', '30')
    add_list_item('HÌNH 4.3: Giao diện đăng nhập tài khoản', '31')
    add_list_item('HÌNH 4.4: Bảng điều khiển trang bệnh nhân', '31')
    add_list_item('HÌNH 4.5: Giao diện đánh giá chất lượng sinh viên', '32')
    add_list_item('HÌNH 4.6: Lịch sử nhận việc của bệnh nhân', '32')
    add_list_item('HÌNH 4.7: Bảng điều khiển trang sinh viên y khoa', '33')
    add_list_item('HÌNH 4.8: Form cập nhật thông tin cá nhân sinh viên', '33')
    add_list_item('HÌNH 4.9: Form gửi hồ sơ xin xác thực tài khoản sinh viên', '34')
    add_list_item('HÌNH 4.10: Form xin cấp quyền đăng tin ứng tuyển', '34')
    add_list_item('HÌNH 4.11: Lịch sử nhận việc của sinh viên', '35')
    add_list_item('HÌNH 4.12: Giao diện hội thoại chat thời gian thực', '35')
    add_list_item('HÌNH 4.13: Giao diện quản lý danh sách bạn bè', '35')
    add_list_item('HÌNH 4.14: Giao diện quản lý các tin đăng yêu thích', '36')
    add_list_item('HÌNH 4.15: Giao diện hỗ trợ tài khoản người dùng', '36')
    add_list_item('HÌNH 4.16: Bảng điều khiển trang quản trị (Admin Dashboard)', '37')
    add_list_item('HÌNH 4.17: Trang quản lý người dùng của admin', '37')
    add_list_item('HÌNH 4.18: Trang kiểm duyệt và quản lý bài đăng', '37')
    add_list_item('HÌNH 4.19: Trang quản trị tin tuyển dụng', '38')
    add_list_item('HÌNH 4.20: Trang quản trị tin ứng tuyển', '38')
    add_list_item('HÌNH 4.21: Giao diện admin gửi thông báo hệ thống', '38')

    for _ in range(2):
        doc.add_paragraph()
    p_note = doc.add_paragraph()
    r = p_note.add_run('Ghi chú:\n')
    r.bold = True
    r.italic = True
    r.font.size = Pt(10)
    r = p_note.add_run('- Xếp sau trang Mục lục\n'
                       '- Chữ số thứ nhất chỉ tên chương\n'
                       '- Chữ số thứ hai chỉ thứ tự bảng biểu, sơ đồ, hình... trong mỗi chương\n'
                       '- Ở cuối mỗi bảng biểu, sơ đồ, hình... trong mỗi chương phải có ghi chú, giải thích, nêu rõ nguồn trích hoặc sao chụp,...')
    r.italic = True
    r.font.size = Pt(10)

    doc.add_page_break()

    # ====== PAGE 15: KÍ HIỆU CÁC CỤM TỪ VIẾT TẮT ======
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('(Mẫu)\n')
    r.italic = True
    r.font.size = Pt(10)
    r = p.add_run('KÍ HIỆU CÁC CỤM TỪ VIẾT TẮT\n')
    r.bold = True
    r.font.size = Pt(14)
    r = p.add_run('(Được xếp sau trang Danh mục các bảng, sơ đồ, hình)')
    r.italic = True
    r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(24)

    # Standard two columns alignment
    abbr_table = doc.add_table(rows=11, cols=2)
    abbr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    abbr_table.rows[0].cells[0].width = Cm(4.0)
    abbr_table.rows[0].cells[1].width = Cm(11.0)
    
    abbrs = [
        ('SV', 'Sinh viên'),
        ('BN', 'Bệnh nhân'),
        ('GVHD', 'Giảng viên hướng dẫn'),
        ('CQTT', 'Cơ quan thực tập'),
        ('CSDL', 'Cơ sở dữ liệu'),
        ('PHP', 'Hypertext Preprocessor (Ngôn ngữ lập trình kịch bản phía máy chủ)'),
        ('SQL', 'Structured Query Language (Ngôn ngữ truy vấn mang tính cấu trúc)'),
        ('HTML', 'HyperText Markup Language (Ngôn ngữ đánh dấu siêu văn bản)'),
        ('CSS', 'Cascading Style Sheets (Tờ mã phong cách dạng thác nước)'),
        ('ERD', 'Entity Relationship Diagram (Biểu đồ quan hệ thực thể)'),
        ('UML', 'Unified Modeling Language (Ngôn ngữ mô hình hóa thống nhất)')
    ]
    
    for i, (abbr, desc) in enumerate(abbrs):
        row = abbr_table.rows[i]
        
        # Abbr
        p_abbr = row.cells[0].paragraphs[0]
        r_abbr = p_abbr.add_run(abbr + ' :')
        r_abbr.bold = True
        r_abbr.font.size = Pt(13)
        r_abbr.font.name = 'Times New Roman'
        
        # Desc
        p_desc = row.cells[1].paragraphs[0]
        r_desc = p_desc.add_run(desc)
        r_desc.font.size = Pt(13)
        r_desc.font.name = 'Times New Roman'

    # Clear borders of abbr table
    for row in abbr_table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for b in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{b}')
                border.set(qn('w:val'), 'none')
                tcBorders.append(border)
            tcPr.append(tcBorders)

    for _ in range(3):
        doc.add_paragraph()
    p_note = doc.add_paragraph()
    r = p_note.add_run('Ghi chú:\n')
    r.bold = True
    r.italic = True
    r.font.size = Pt(10)
    r = p_note.add_run('Cụm từ viết tắt là các chữ cái và các ký hiệu thay chữ được viết liền nhau, để thay cho một cụm từ có nghĩa thông thường được lặp nhiều lần trong văn bản hoặc được mọi người mặc nhiên chấp nhận.')
    r.italic = True
    r.font.size = Pt(10)

    doc.add_page_break()

    # ====== PARSE AND APPEND BODY FROM DACN.tex ======
    print("Parsing LaTeX file and appending body...")
    parse_latex_to_docx(doc, 'DACN.tex')
    
    # Save document
    output_filename = 'Bao_cao_tot_nghiep.docx'
    try:
        doc.save(output_filename)
        print(f"Báo cáo tốt nghiệp đã được tạo thành công: {output_filename}")
    except PermissionError:
        output_filename = 'Bao_cao_tot_nghiep_v2.docx'
        doc.save(output_filename)
        print(f"Lưu ý: Tệp Bao_cao_tot_nghiep.docx đang mở trong Word, do đó báo cáo đã được lưu thành: {output_filename}")


def parse_latex_to_docx(doc, tex_filepath):
    if not os.path.exists(tex_filepath):
        print(f"Error: LaTeX file not found at {tex_filepath}")
        return
        
    with open(tex_filepath, 'r', encoding='utf-8') as f:
        content = f.read()
        
    # Find start of document content
    doc_start_match = re.search(r'\\begin\{document\}', content)
    if not doc_start_match:
        print("Error: \\begin{document} not found in LaTeX file.")
        return
        
    body_content = content[doc_start_match.end():]
    
    # Strip some specific document template structures (like title page, signature pages, acknowledgments)
    # since we have already hardcoded them in their exact template layout!
    # Let's search for the first actual chapter or section of the thesis body.
    # The thesis body starts with \chapter{TỔNG QUAN} or similar.
    # Let's search for the first chapter.
    first_chapter_match = re.search(r'\\chapter', body_content)
    if first_chapter_match:
        # Keep everything from the first chapter onwards
        body_content = body_content[first_chapter_match.start():]
    else:
        # Fallback: search for TÓM TẮT ĐỒ ÁN
        tom_tat_match = re.search(r'TÓM TẮT ĐỒ ÁN', body_content)
        if tom_tat_match:
            body_content = body_content[tom_tat_match.start() - 50:] # approximate start

    # Split contents into lines/tokens
    lines = body_content.split('\n')
    
    # Simple state machine to parse latex elements
    in_itemize = False
    in_enumerate = False
    in_table = False
    in_tabular = False
    in_figure = False
    
    table_rows = []
    table_caption = ""
    figure_filename = ""
    figure_caption = ""
    
    item_num = 0
    
    # Dynamic image mapping for typos and differences between LaTeX names and actual file names
    image_mappings = {
        'Giao diện trang chủ.png': 'images/Giao diện trang chủ.png',
        'Giao diện đăng kí.png': 'images/Giao diện đăng kí.png',
        'Giao diện đăng nhập.png': 'images/Giao diện đăng nhập.png',
        'Giao diện bảng điều khiển trang bệnh nhân.png': 'images/Giao diện bảng điều khiển trang bệnh nhân.png',
        'Giao diện đánh giá trang bệnh nhân.png': 'images/Giao diện đánh giá trang bệnh nhân.png',
        'Giao diện lịch sử nhận việc trang bệnh nhân.png': 'images/Giao diện lịch sử nhận việc trang bệnh nhân.png',
        'Giao diện bảng điều khiển trang sinh viên.png': 'images/Giao diện bảng điều khiển trang sinh viên.png',
        'Giao diện form cập nhật hồ sơ trang sinh viên.png': 'images/Giao diện form cập nhật hồ sơ trang sinh viên.png',
        'Giao diện form xin xác thực tài khoản.png': 'images/Giao diện form xin xác thực tài khoản.png',
        'Giao diện form xin cấp quyền đăng tin.png': 'images/Giao diện form xin cấp quyền đăng tin.png',
        'Giao diện lịch sử nhận việc trang sinh viên.png': 'images/Giao diện lịch sử nhận việc trang sinh viên.png',
        'Giao diện hội thoại.png': 'images/Giao diện hội thoại.png',
        'Giao diện trang bạn bè.png': 'images/Giao diện trang bạn bè.png',
        'Giao diện trang bài tuyển yêu thích.png': 'images/Giao diện trang bài tuyển yêu thích.png',
        'Giao diện trang hỗ trợ tài khoản.png': 'images/Giao diện trang hỗ trợ tài khoản.png',
        'Giao diện bảng điều khiển trang quản trị.png': 'images/Giao diện bảng điều khiển trang quản trị.png',
        'Giao diện quản lý người dùng.png': 'images/Giao diện quản lý người dùng.png',
        'Giao diện quản lý bài viết.png': 'images/Giao diện quản lý bài viết.png',
        'Giao diện đăng tin tuyển trang quản trị.png': 'images/admin_tintuyen.png',
        'Giao diện đăng tin ứng tuyển trang quản trị.png': 'images/admin_tinungtuyen.png',
        'Giao dện gửi thông báo .png': 'images/admin_guithongbao.png',
        'usecase.png': 'images_doan/usecase.jpg',
        'ERD.png': 'images_doan/ERD.png'
    }

    # Clean text helper
    def clean_latex_markup(text):
        # Remove LaTeX comments
        text = re.sub(r'(?<!\\)%.*', '', text)
        # Replace inline formatting
        text = re.sub(r'\\texttt\{([^}]+)\}', r'\1', text)
        text = re.sub(r'\\textbf\{([^}]+)\}', r'\1', text)
        text = re.sub(r'\\textit\{([^}]+)\}', r'\1', text)
        # Replace mathematical/symbol characters
        text = text.replace(r'\rightarrow', '→')
        text = text.replace(r'\dots', '...')
        text = text.replace(r'\ldots', '...')
        text = text.replace(r'\_', '_')
        text = text.replace(r'\&', '&')
        text = text.replace(r'\$', '$')
        text = text.replace(r'\{', '{')
        text = text.replace(r'\}', '}')
        # Strip outer spaces
        return text.strip()

    def process_inline_formatting(p, raw_text):
        # We parse the text and add runs for bold, italic and monospace text
        # Simple parser for \textbf, \textit, \texttt
        pattern = re.compile(r'\\(textbf|textit|texttt)\{([^}]+)\}')
        
        last_end = 0
        for match in pattern.finditer(raw_text):
            start, end = match.span()
            cmd = match.group(1)
            inner_text = match.group(2)
            
            # Normal text before match
            if start > last_end:
                normal_txt = clean_latex_markup(raw_text[last_end:start])
                if normal_txt:
                    r = p.add_run(normal_txt)
                    r.font.name = 'Times New Roman'
                    
            # Styled text
            styled_txt = clean_latex_markup(inner_text)
            r = p.add_run(styled_txt)
            r.font.name = 'Times New Roman'
            if cmd == 'textbf':
                r.bold = True
            elif cmd == 'textit':
                r.italic = True
            elif cmd == 'texttt':
                r.font.name = 'Courier New'
                r.font.size = Pt(11)
                
            last_end = end
            
        if last_end < len(raw_text):
            rem_txt = clean_latex_markup(raw_text[last_end:])
            if rem_txt:
                r = p.add_run(rem_txt)
                r.font.name = 'Times New Roman'

    injection_state = {
        'inject_bootstrap': False,
        'inject_ai_guide': False,
        'inject_usecase': False,
        'inject_db': False,
        'inject_evaluation': False,
        'inject_appendix': False
    }

    def perform_deferred_injections(doc):
        if injection_state['inject_bootstrap']:
            parse_txt_to_docx(doc, 'vai_tro_bootstrap.txt')
            injection_state['inject_bootstrap'] = False
            
        if injection_state['inject_ai_guide']:
            parse_markdown_to_docx(doc, 'HUONG_DAN_AI.md')
            injection_state['inject_ai_guide'] = False
            
        if injection_state['inject_usecase']:
            parse_markdown_to_docx(doc, 'QUY_TRINH_KET_NOI_Y_TE.md')
            injection_state['inject_usecase'] = False
            
        if injection_state['inject_db']:
            parse_markdown_to_docx(doc, 'phan_tich_powerdesigner.md')
            parse_markdown_to_docx(doc, 'phan_tich_kieu_du_lieu.md')
            injection_state['inject_db'] = False
            
        if injection_state['inject_evaluation']:
            parse_markdown_to_docx(doc, 'DANH_GIA_QUY_TRINH_KET_NOI_Y_TE.md')
            parse_markdown_to_docx(doc, 'HUONG_DAN_TEST_POSTMAN.md')
            parse_txt_to_docx(doc, 'uu_diem_han_che.txt')
            injection_state['inject_evaluation'] = False
            
        if injection_state['inject_appendix']:
            parse_markdown_to_docx(doc, 'HUONG_DAN_OVERLEAF.md')
            parse_markdown_to_docx(doc, 'HUONG_DAN_FIX_NGROK.md')
            injection_state['inject_appendix'] = False

    idx = 0
    while idx < len(lines):
        line = lines[idx].strip()
        
        # Skip empty lines
        if not line:
            idx += 1
            continue
            
        # Ignore end document
        if '\\end{document}' in line:
            perform_deferred_injections(doc)
            break
            
        # CHAPTER
        if line.startswith('\\chapter{'):
            perform_deferred_injections(doc)
            chapter_name = re.match(r'\\chapter\{([^}]+)\}', line).group(1)
            # Add Page Break before every chapter except the first one
            p_ch = doc.add_paragraph()
            p_ch.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_ch.paragraph_format.space_before = Pt(24)
            p_ch.paragraph_format.space_after = Pt(18)
            p_ch.text = "" # Clear
            r = p_ch.add_run(f'{chapter_name.upper()}')
            r.bold = True
            r.font.size = Pt(14)
            r.font.name = 'Times New Roman'
            doc.add_paragraph() # Spacing
            idx += 1
            continue
            
        # SECTION
        elif line.startswith('\\section{') or line.startswith('\\section*{'):
            perform_deferred_injections(doc)
            section_name = re.match(r'\\section\*?\{([^}]+)\}', line).group(1)
            p_sec = doc.add_paragraph()
            p_sec.paragraph_format.space_before = Pt(12)
            p_sec.paragraph_format.space_after = Pt(6)
            p_sec.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p_sec.add_run(section_name)
            r.bold = True
            r.font.size = Pt(13)
            r.font.name = 'Times New Roman'
            
            # Check for triggers
            if 'Phương pháp phân tích thiết kế hệ thống' in section_name:
                injection_state['inject_ai_guide'] = True
            elif 'Phụ lục A' in section_name:
                injection_state['inject_appendix'] = True
                
            idx += 1
            continue
            
        # SUBSECTION
        elif line.startswith('\\subsection{') or line.startswith('\\subsection*{'):
            perform_deferred_injections(doc)
            subsec_name = re.match(r'\\subsection\*?\{([^}]+)\}', line).group(1)
            p_sub = doc.add_paragraph()
            p_sub.paragraph_format.space_before = Pt(8)
            p_sub.paragraph_format.space_after = Pt(4)
            p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p_sub.add_run(subsec_name)
            r.bold = True
            r.italic = True
            r.font.size = Pt(13)
            r.font.name = 'Times New Roman'
            
            # Check for triggers
            if 'Bootstrap Framework' in subsec_name:
                injection_state['inject_bootstrap'] = True
            elif 'Biểu đồ Use Case' in subsec_name:
                injection_state['inject_usecase'] = True
            elif 'Hạn chế' in subsec_name:
                injection_state['inject_evaluation'] = True
                
            idx += 1
            continue
            
        # SUBSUBSECTION
        elif line.startswith('\\subsubsection{'):
            perform_deferred_injections(doc)
            subsubsec_name = re.match(r'\\subsubsection\{([^}]+)\}', line).group(1)
            p_ssub = doc.add_paragraph()
            p_ssub.paragraph_format.space_before = Pt(6)
            p_ssub.paragraph_format.space_after = Pt(3)
            p_ssub.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p_ssub.add_run(subsubsec_name)
            r.bold = True
            r.font.size = Pt(13)
            r.font.name = 'Times New Roman'
            
            # Check for triggers
            if 'Mô tả các bảng dữ liệu' in subsubsec_name:
                injection_state['inject_db'] = True
                
            idx += 1
            continue

        # ITEMIZE LISTS
        elif line.startswith('\\begin{itemize}'):
            in_itemize = True
            idx += 1
            continue
        elif line.startswith('\\end{itemize}'):
            in_itemize = False
            idx += 1
            continue
            
        # ENUMERATE LISTS
        elif line.startswith('\\begin{enumerate}'):
            in_enumerate = True
            item_num = 0
            idx += 1
            continue
        elif line.startswith('\\end{enumerate}'):
            in_enumerate = False
            idx += 1
            continue
            
        # ITEMS
        elif line.startswith('\\item'):
            item_text = line[5:].strip()
            # If the item spans multiple lines, we can check next lines, but simple check first
            p_item = doc.add_paragraph()
            p_item.paragraph_format.space_after = Pt(4)
            p_item.paragraph_format.left_indent = Cm(0.8)
            
            if in_itemize:
                p_item.add_run('•  ').font.name = 'Times New Roman'
            elif in_enumerate:
                item_num += 1
                p_item.add_run(f'{item_num}.  ').font.name = 'Times New Roman'
                
            process_inline_formatting(p_item, item_text)
            idx += 1
            continue
            
        # FIGURES
        elif line.startswith('\\begin{figure}'):
            in_figure = True
            figure_filename = ""
            figure_caption = ""
            idx += 1
            continue
        elif line.startswith('\\end{figure}'):
            in_figure = False
            if figure_filename:
                # Add picture
                p_fig = doc.add_paragraph()
                p_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_fig.paragraph_format.space_before = Pt(6)
                p_fig.paragraph_format.space_after = Pt(6)
                
                # Resolve filename
                actual_filename = figure_filename
                # Strip curly brackets and extensions
                actual_filename = actual_filename.replace('{', '').replace('}', '')
                if not actual_filename.endswith('.png') and not actual_filename.endswith('.jpg'):
                    actual_filename += '.png'
                    
                mapped_path = image_mappings.get(actual_filename, os.path.join('images', actual_filename))
                
                # Fallback check
                if not os.path.exists(mapped_path):
                    # Check in images_doan
                    fallback_path = os.path.join('images_doan', actual_filename)
                    if os.path.exists(fallback_path):
                        mapped_path = fallback_path
                    else:
                        # Check in Ảnh Đồ Án
                        fallback_path = os.path.join('Ảnh Đồ Án', actual_filename)
                        if os.path.exists(fallback_path):
                            mapped_path = fallback_path
                
                print(f"Inserting image: {mapped_path} for LaTeX ref: {figure_filename}")
                
                mapped_path = get_safe_image_path(mapped_path)
                if mapped_path:
                    try:
                        p_fig.add_run().add_picture(mapped_path, width=Cm(14.0))
                    except Exception as e:
                        p_fig.add_run(f'[Lỗi hiển thị hình ảnh {actual_filename}: {str(e)}]').italic = True
                else:
                    p_fig.add_run(f'[Hình ảnh minh họa: {actual_filename}]').italic = True
                    print(f"Warning: Image file not found: {mapped_path}")
                
                # Add caption
                if figure_caption:
                    p_cap = doc.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.paragraph_format.space_after = Pt(12)
                    r = p_cap.add_run(f'Hình: {clean_latex_markup(figure_caption)}')
                    r.italic = True
                    r.font.size = Pt(11)
                    r.font.name = 'Times New Roman'
                    
            idx += 1
            continue
            
        elif in_figure and line.startswith('\\includegraphics'):
            # Extract filename
            match = re.search(r'\\includegraphics\[[^\]]+\]\{([^}]+)\}', line)
            if not match:
                match = re.search(r'\\includegraphics\{([^}]+)\}', line)
            if match:
                figure_filename = match.group(1)
            idx += 1
            continue
            
        elif in_figure and line.startswith('\\caption{'):
            figure_caption = re.match(r'\\caption\{([^}]+)\}', line).group(1)
            idx += 1
            continue

        # TABLES
        elif line.startswith('\\begin{table}'):
            in_table = True
            table_caption = ""
            idx += 1
            continue
        elif line.startswith('\\end{table}'):
            in_table = False
            idx += 1
            continue
            
        elif in_table and line.startswith('\\caption{'):
            table_caption = re.match(r'\\caption\{([^}]+)\}', line).group(1)
            idx += 1
            continue
            
        elif line.startswith('\\begin{tabular}'):
            in_tabular = True
            table_rows = []
            # Parse column alignment to know column count
            col_spec = re.match(r'\\begin\{tabular\}\{([^}]+)\}', line).group(1)
            # Count character elements that represent columns (l, c, r, p)
            col_count = len(re.findall(r'[lcrp]', col_spec))
            idx += 1
            continue
            
        elif line.startswith('\\end{tabular}'):
            in_tabular = False
            
            # Draw table in Word
            if table_rows:
                # Table title
                if table_caption:
                    p_cap = doc.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.paragraph_format.space_before = Pt(12)
                    p_cap.paragraph_format.space_after = Pt(6)
                    r = p_cap.add_run(f'Bảng: {clean_latex_markup(table_caption)}')
                    r.bold = True
                    r.font.size = Pt(11)
                    r.font.name = 'Times New Roman'
                
                rows_count = len(table_rows)
                cols_count = len(table_rows[0])
                
                word_table = doc.add_table(rows=rows_count, cols=cols_count)
                word_table.style = 'Table Grid'
                word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                for r_idx, r_data in enumerate(table_rows):
                    for c_idx, cell_text in enumerate(r_data):
                        cell = word_table.cell(r_idx, c_idx)
                        p_cell = cell.paragraphs[0]
                        p_cell.paragraph_format.space_before = Pt(2)
                        p_cell.paragraph_format.space_after = Pt(2)
                        
                        # Heading formatting for first row
                        if r_idx == 0:
                            p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            r_run = p_cell.add_run(clean_latex_markup(cell_text))
                            r_run.bold = True
                        else:
                            p_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            r_run = p_cell.add_run(clean_latex_markup(cell_text))
                        
                        r_run.font.name = 'Times New Roman'
                        r_run.font.size = Pt(11)
                        
                doc.add_paragraph() # Spacing after table
            idx += 1
            continue
            
        elif in_tabular:
            # Handle tabular row data
            # Ignore borders like \hline
            if line.startswith('\\hline'):
                idx += 1
                continue
            # Parse row: separated by & and ending with \\
            row_clean = line.replace('\\\\', '').strip()
            if row_clean:
                parts = [p.strip() for p in row_clean.split('&')]
                table_rows.append(parts)
            idx += 1
            continue
            
        # OTHER COMMANDS TO IGNORE
        elif line.startswith('\\addcontentsline') or line.startswith('\\thispagestyle') or line.startswith('\\newpage') or line.startswith('\\noindent') or line.startswith('\\vspace') or line.startswith('\\rule') or line.startswith('\\begin') or line.startswith('\\end'):
            idx += 1
            continue
            
        # NORMAL TEXT PARAGRAPH
        else:
            # Join consecutive text lines until next command or empty line
            p_text = line
            next_idx = idx + 1
            while next_idx < len(lines):
                next_line = lines[next_idx].strip()
                if not next_line or next_line.startswith('\\') or next_line.startswith('%'):
                    break
                p_text += " " + next_line
                next_idx = next_idx + 1
                
            idx = next_idx
            
            p_normal = doc.add_paragraph()
            p_normal.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_normal.paragraph_format.line_spacing = 1.5
            p_normal.paragraph_format.first_line_indent = Cm(1.25)
            p_normal.paragraph_format.space_after = Pt(6)
            
            process_inline_formatting(p_normal, p_text)
            continue

if __name__ == '__main__':
    create_report()
