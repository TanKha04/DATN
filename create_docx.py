from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(13)

# ====== HEADER TABLE (2 columns) ======
header_table = doc.add_table(rows=2, cols=2)
header_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Left column
c0 = header_table.cell(0, 0)
p = c0.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('TRƯỜNG KỸ THUẬT VÀ CÔNG NGHỆ')
r.bold = True
r.font.size = Pt(11)
r.font.name = 'Times New Roman'

c1 = header_table.cell(1, 0)
p = c1.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('KHOA CÔNG NGHỆ THÔNG TIN')
r.bold = True
r.font.size = Pt(11)
r.font.name = 'Times New Roman'

# Right column
c2 = header_table.cell(0, 1)
p = c2.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM')
r.bold = True
r.font.size = Pt(12)
r.font.name = 'Times New Roman'

c3 = header_table.cell(1, 1)
p = c3.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Độc lập – Tự do – Hạnh phúc')
r.italic = True
r.font.size = Pt(11)
r.font.name = 'Times New Roman'

# Remove borders from header table
for row in header_table.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is not None:
            tcPr.remove(tcBorders)
        borders = '<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/></w:tcBorders>'
        from lxml import etree
        tcPr.append(etree.fromstring(borders))

doc.add_paragraph()

# ====== TITLE ======
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('ĐỀ CƯƠNG CHI TIẾT')
r.bold = True
r.font.size = Pt(15)
r.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('ĐỒ ÁN TỐT NGHIỆP NGÀNH CÔNG NGHỆ THÔNG TIN')
r.bold = True
r.font.size = Pt(14)
r.font.name = 'Times New Roman'

doc.add_paragraph()

# ====== STUDENT INFO ======
def add_info_line(label, value, label2=None, value2=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(label)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)
    r = p.add_run(value)
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)
    if label2:
        r = p.add_run('     ' + label2)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)
        r = p.add_run(value2)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)

add_info_line('Họ tên sinh viên: ', 'Trầm Tấn Khả', 'MSSV: ', '110122026')
add_info_line('Lớp: ', 'DA22TTD', 'Khóa: ', '2022')

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run('Tên đề tài: ')
r.font.name = 'Times New Roman'
r.font.size = Pt(13)
r = p.add_run('Kết nối sinh viên ngành y với người có nhu cầu chăm sóc sức khỏe tại nhà')
r.font.name = 'Times New Roman'
r.font.size = Pt(13)

doc.add_paragraph()

# ====== HELPER FUNCTIONS ======
def add_section_title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)

def add_bullet(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1)
    r = p.add_run('- ' + text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

def add_chapter(title, content):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1)
    r = p.add_run(title)
    r.italic = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1.5)
    r = p.add_run(content)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

# ====== 1. MỤC TIÊU ======
add_section_title('1. Mục tiêu của đề tài:')
add_bullet('Xây dựng nền tảng web kết nối sinh viên y khoa với người dân có nhu cầu chăm sóc sức khỏe tại nhà.')
add_bullet('Tạo môi trường an toàn, minh bạch để bệnh nhân tìm hỗ trợ y tế và sinh viên tìm cơ hội thực hành lâm sàng.')
add_bullet('Tích hợp trí tuệ nhân tạo (AI) hỗ trợ tư vấn sức khỏe sơ bộ và gợi ý sinh viên phù hợp.')
add_bullet('Xây dựng hệ thống xác minh danh tính sinh viên đảm bảo độ tin cậy cho người sử dụng.')

# ====== 2. NỘI DUNG ======
add_section_title('2. Nội dung thực hiện:')
add_bullet('Phân tích yêu cầu hệ thống với 3 vai trò: Bệnh nhân, Sinh viên Y khoa, Quản trị viên.')
add_bullet('Thiết kế cơ sở dữ liệu MySQL (16 bảng: users, posts, messages, ratings, verifications, ...).')
add_bullet('Xây dựng các module: đăng ký/đăng nhập, đăng tin tuyển dụng/ứng tuyển, tin nhắn, gọi video WebRTC.')
add_bullet('Tích hợp Trợ lý AI (Google Gemini): tư vấn sức khỏe, kiểm tra triệu chứng, tìm sinh viên phù hợp.')
add_bullet('Xây dựng hệ thống xác minh sinh viên, đánh giá, kết bạn, thông báo, bảng xếp hạng.')
add_bullet('Xây dựng trang quản trị viên: quản lý người dùng, bài đăng, xác minh, báo cáo vi phạm.')

# ====== 3. PHƯƠNG PHÁP ======
add_section_title('3. Phương pháp thực hiện:')
add_bullet('Nghiên cứu tài liệu về PHP, MySQL, Bootstrap, REST API, WebRTC, Google Gemini AI.')
add_bullet('Phân tích thiết kế hệ thống bằng UML: Use Case Diagram, ERD, Class Diagram.')
add_bullet('Phát triển theo mô hình Agile, kiểm thử API bằng Postman, triển khai bằng Docker.')

# ====== 4. BỐ CỤC ======
add_section_title('4. Bố cục đề tài:')
add_chapter('Chương 1: Tổng quan đề tài',
    'Giới thiệu bối cảnh thực tiễn, lý do chọn đề tài, mục tiêu, phạm vi nghiên cứu, ý nghĩa khoa học và thực tiễn của đề tài.')
add_chapter('Chương 2: Cơ sở lý thuyết',
    'Trình bày lý thuyết về PHP, MySQL, Bootstrap 5, REST API, WebRTC (gọi video P2P), Google Gemini AI API, xác thực email SMTP, OAuth (Facebook Login), Docker.')
add_chapter('Chương 3: Phân tích và thiết kế hệ thống',
    'Phân tích yêu cầu chức năng và phi chức năng. Thiết kế Use Case Diagram cho 3 actor. Thiết kế ERD (16 bảng), Class Diagram, thiết kế giao diện người dùng (wireframe).')
add_chapter('Chương 4: Xây dựng và triển khai',
    'Cài đặt môi trường, xây dựng các module: đăng ký/đăng nhập, bài đăng, tin nhắn, gọi video WebRTC, trợ lý AI Gemini, xác minh sinh viên, dashboard. Triển khai hosting.')
add_chapter('Chương 5: Thực nghiệm và đánh giá',
    'Kiểm thử các chức năng chính, kiểm thử API bằng Postman. Đánh giá hiệu suất, bảo mật. So sánh với các giải pháp tương tự, khảo sát trải nghiệm người dùng.')
add_chapter('Chương 6: Kết luận và hướng phát triển',
    'Tổng kết kết quả đạt được, hạn chế còn tồn tại. Đề xuất hướng phát triển: tích hợp thanh toán, ứng dụng mobile, nâng cấp AI, lịch hẹn tự động, mở rộng phạm vi địa lý.')

# ====== 5. TÀI LIỆU THAM KHẢO ======
add_section_title('5. Tài liệu tham khảo:')
refs = [
    '[1] PHP Documentation - https://www.php.net/docs.php',
    '[2] MySQL 8.0 Reference Manual - https://dev.mysql.com/doc/refman/8.0/en/',
    '[3] Bootstrap 5 Documentation - https://getbootstrap.com/docs/5.3/',
    '[4] WebRTC API - MDN Web Docs - https://developer.mozilla.org/en-US/docs/Web/API/WebRTC_API',
    '[5] Google Gemini AI API - https://ai.google.dev/docs',
    '[6] Docker Documentation - https://docs.docker.com/',
    '[7] RESTful API Design Best Practices - https://restfulapi.net/',
    '[8] PHPMailer - https://github.com/PHPMailer/PHPMailer',
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1)
    r = p.add_run(ref)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

# ====== 6. KẾ HOẠCH ======
add_section_title('6. Kế hoạch thực hiện đề tài:')

# Table data
schedule = [
    ('1', '20/4 – 26/4/2026', 'Tìm hiểu đề tài, khảo sát yêu cầu, thu thập tài liệu tham khảo', 'Chọn tài liệu'),
    ('2', '27/4 – 3/5/2026', 'Nghiên cứu lý thuyết: PHP, MySQL, Bootstrap, REST API, WebRTC, Gemini AI', 'Đọc tài liệu'),
    ('3', '4/5 – 10/5/2026', 'Phân tích yêu cầu, thiết kế Use Case, ERD (16 bảng), thiết kế giao diện', 'Vẽ sơ đồ UML'),
    ('4', '11/5 – 17/5/2026', 'Xây dựng CSDL, module đăng ký/đăng nhập, xác thực email, phân quyền vai trò', 'Tạo database'),
    ('5', '18/5 – 24/5/2026', 'Xây dựng module đăng tin tuyển dụng/ứng tuyển, tìm kiếm, lọc, xem chi tiết', 'Code frontend'),
    ('6', '25/5 – 31/5/2026', 'Xây dựng module tin nhắn, hội thoại, gọi video WebRTC, kết bạn, thông báo', 'Code backend'),
    ('7', '1/6 – 7/6/2026', 'Tích hợp Trợ lý AI Gemini, kiểm tra triệu chứng, xác minh SV, đánh giá', 'Tích hợp AI'),
    ('8', '8/6 – 14/6/2026', 'Xây dựng Dashboard (Bệnh nhân, Sinh viên, Admin), Docker hóa, triển khai', 'Deploy hosting'),
    ('9', '15/6 – 21/6/2026', 'Kiểm thử hệ thống, kiểm thử API (Postman), sửa lỗi, tối ưu hiệu suất', 'Test Postman'),
    ('10', '22/6 – 28/6/2026', 'Hoàn thiện báo cáo đồ án, chuẩn bị slide, bảo vệ đồ án', 'Viết báo cáo'),
]

table = doc.add_table(rows=11, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
headers = ['Tuần', 'Từ ngày - đến ngày', 'Công việc thực hiện', 'Ghi chú']
for i, h in enumerate(headers):
    cell = table.cell(0, i)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)

# Data rows
for idx, (week, dates, task, note) in enumerate(schedule):
    row = table.rows[idx + 1]
    
    # Tuần
    p = row.cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(week)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    
    # Ngày
    p = row.cells[1].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(dates)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    
    # Công việc
    p = row.cells[2].paragraphs[0]
    r = p.add_run(task)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    
    # Ghi chú
    p = row.cells[3].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(note)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)

# Column widths
for row in table.rows:
    row.cells[0].width = Cm(1.5)
    row.cells[1].width = Cm(4)
    row.cells[2].width = Cm(9)
    row.cells[3].width = Cm(3)

doc.add_paragraph()
doc.add_paragraph()

# ====== SIGNATURES ======
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run('Vĩnh Long, ngày       tháng       năm 2026')
r.italic = True
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

doc.add_paragraph()

sig_table = doc.add_table(rows=3, cols=2)
sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Remove borders
for row in sig_table.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        borders = '<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/></w:tcBorders>'
        from lxml import etree
        tcPr.append(etree.fromstring(borders))

# Row 0: Titles
for i, title in enumerate(['GIẢNG VIÊN HƯỚNG DẪN', 'SINH VIÊN THỰC HIỆN']):
    p = sig_table.cell(0, i).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)

# Row 1: Empty for signature space
for i in range(2):
    p = sig_table.cell(1, i).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('\n\n')
    r.font.size = Pt(13)

# Row 2: Names
for i, name in enumerate(['Nguyên Hoàng Duy Thiện', 'Trầm Tấn Khả']):
    p = sig_table.cell(2, i).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(name)
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)

# Save
doc.save(r'c:\xampp\htdocs\DACN2\de_cuong_chi_tiet.docx')
print('DONE! File saved: de_cuong_chi_tiet.docx')
